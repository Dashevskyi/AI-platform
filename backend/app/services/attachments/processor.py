"""
Attachment processing pipeline:
1. Extract text from file (PDF, DOCX, XLSX, text, image via OCR, audio via Whisper)
2. Generate summary via LLM
3. Chunk text
4. Generate embeddings for each chunk
"""
import asyncio
import base64
import io
import logging
import uuid

from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.message_attachment import MessageAttachment
from app.models.message_attachment_chunk import MessageAttachmentChunk
from app.models.tenant_shell_config import TenantShellConfig
from app.providers.factory import get_provider
from app.services.storage import read_file
from app.services.kb.embedder import chunk_text
from app.core.config import settings
from app.core.database import async_session

logger = logging.getLogger(__name__)


# ============================================================
# Content extraction by file type
# ============================================================

def extract_text_content(file_bytes: bytes, filename: str) -> str:
    """Extract text from plain text files."""
    return file_bytes.decode("utf-8", errors="replace")


def extract_pdf_content(file_bytes: bytes) -> str:
    """Extract text from PDF — text-layer only. Fast path. Returns possibly
    sparse/empty string on scanned PDFs (no text layer). See
    `extract_pdf_content_async` for the OCR-fallback variant."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


async def _ocr_image_bytes(png_bytes: bytes, filename: str) -> str:
    """Send a single rendered page to the dual-pass OCR endpoint. Returns
    empty string on failure — caller decides what to do."""
    url = settings.OCR_URL
    if not url:
        return ""
    import httpx
    try:
        async with httpx.AsyncClient(timeout=settings.OCR_TIMEOUT_SECONDS) as client:
            resp = await client.post(
                url,
                files={"file": (filename, png_bytes, "image/png")},
                data={"lang": "cyrillic"},  # ignored by /auto
            )
            resp.raise_for_status()
            return (resp.json().get("text") or "").strip()
    except Exception as e:
        logger.warning(f"PDF page OCR failed for {filename}: {type(e).__name__}: {e!r}")
        return ""


async def extract_pdf_content_async(file_bytes: bytes, filename: str = "document.pdf") -> str:
    """Per-page PDF extraction with OCR fallback for scanned pages.

    Strategy: for each page, take pypdf's text-layer output. If shorter than
    PDF_PAGE_TEXT_LAYER_MIN_CHARS, treat the page as a scan, render it to PNG
    via PyMuPDF and run /v1/ocr/auto on it. Capped at PDF_OCR_MAX_PAGES OCR
    operations so a 500-page scan can't lock the server."""
    import asyncio
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    page_texts: dict[int, str] = {}
    ocr_targets: list[int] = []
    threshold = settings.PDF_PAGE_TEXT_LAYER_MIN_CHARS
    for idx, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if len(text) >= threshold:
            page_texts[idx] = text
        else:
            ocr_targets.append(idx)

    # Cap how many pages we'll OCR.
    max_ocr = settings.PDF_OCR_MAX_PAGES
    if len(ocr_targets) > max_ocr:
        logger.warning(
            f"PDF {filename}: {len(ocr_targets)} scanned pages, OCR'ing first {max_ocr}"
        )
        ocr_targets = ocr_targets[:max_ocr]

    if ocr_targets:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed — PDF OCR fallback disabled")
            fitz = None  # type: ignore

        if fitz is not None:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            try:
                # Render in a thread pool — PyMuPDF is blocking C code.
                loop = asyncio.get_event_loop()
                dpi = settings.PDF_OCR_RENDER_DPI

                def _render_page(page_idx: int) -> bytes:
                    page = doc.load_page(page_idx)
                    pix = page.get_pixmap(dpi=dpi)
                    return pix.tobytes("png")

                async def _process_page(page_idx: int) -> tuple[int, str]:
                    png_bytes = await loop.run_in_executor(None, _render_page, page_idx)
                    text = await _ocr_image_bytes(png_bytes, f"{filename}.page-{page_idx + 1}.png")
                    return page_idx, text

                import time
                t0 = time.time()
                results = await asyncio.gather(*[_process_page(i) for i in ocr_targets])
                elapsed = time.time() - t0
                ocr_chars = 0
                for page_idx, text in results:
                    if text:
                        page_texts[page_idx] = text
                        ocr_chars += len(text)
                logger.info(
                    f"PDF {filename}: OCR'd {len(ocr_targets)} pages → {ocr_chars} chars in {elapsed:.2f}s"
                )
            finally:
                doc.close()

    # Reassemble in page order.
    ordered = [page_texts[i] for i in sorted(page_texts.keys()) if page_texts[i]]
    return "\n\n".join(ordered)


def extract_docx_content(file_bytes: bytes) -> str:
    """Extract text from DOCX."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def extract_xlsx_content(file_bytes: bytes) -> str:
    """Extract text from XLSX/XLS."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Лист: {sheet_name}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts)


def _extract_image_ocr_tesseract(file_bytes: bytes) -> str:
    """Local Tesseract OCR (CPU). Used as fallback when PaddleOCR server is down."""
    try:
        from PIL import Image
        import pytesseract
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image, lang="rus+eng")
        return text.strip()
    except Exception as e:
        logger.warning(f"Tesseract OCR failed: {e}")
        return ""


async def extract_image_ocr(file_bytes: bytes, filename: str = "image") -> str:
    """Extract text from image via the PaddleOCR GPU server, with Tesseract fallback.
    The GPU server (faster-ocr) handles ru/uk/en via the 'cyrillic' model and takes
    fractions of a second; Tesseract is CPU and slow but covers offline fallback."""
    url = settings.OCR_URL
    if url:
        try:
            import httpx
            ext = (filename.rsplit(".", 1)[-1] or "png").lower()
            mime = {
                "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp",
                "tiff": "image/tiff",
            }.get(ext, "image/png")
            import time
            t0 = time.time()
            async with httpx.AsyncClient(timeout=settings.OCR_TIMEOUT_SECONDS) as client:
                # /v1/ocr/auto ignores `lang`; /v1/ocr accepts it. Send for the
                # single-pass fallback case — server-side it's a no-op for auto.
                resp = await client.post(
                    url,
                    files={"file": (filename, file_bytes, mime)},
                    data={"lang": "cyrillic"},
                )
                resp.raise_for_status()
                data = resp.json()
                text = (data.get("text") or "").strip()
                elapsed = time.time() - t0
                avg_conf = data.get("avg_confidence")
                conf_part = f", avg_conf={avg_conf:.2f}" if isinstance(avg_conf, (int, float)) else ""
                logger.info(
                    f"PaddleOCR {filename} ({len(file_bytes)} bytes) → "
                    f"{data.get('line_count', 0)} lines, {len(text)} chars in {elapsed:.2f}s{conf_part}"
                )
                return text
        except Exception as e:
            logger.warning(
                f"PaddleOCR call failed for {filename}: {type(e).__name__}: {e!r}; "
                f"falling back to Tesseract"
            )
    # Fallback: local CPU OCR
    return _extract_image_ocr_tesseract(file_bytes)


async def describe_image_vision(file_bytes: bytes, filename: str, preferred_model: str | None = None) -> str:
    """
    Describe image using Ollama vision model.
    Priority: preferred_model (if provided and available) → strongest available
    in the order: qwen2-vl > llava:34b > llava:13b > llava > bakllava > llama-vision > moondream.
    """
    vision_model: str | None = None
    try:
        import httpx
        base_url = (settings.OLLAMA_BASE_URL or "http://localhost:11434").rstrip("/")

        # List installed Ollama models
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{base_url}/api/tags")
            installed = [m.get("name", "") for m in resp.json().get("models", [])]

        vision_model: str | None = None

        # 1. Use explicitly preferred model if installed
        if preferred_model:
            for m in installed:
                if m == preferred_model or m.startswith(preferred_model + ":"):
                    vision_model = m
                    break
            if not vision_model:
                logger.warning(f"Preferred vision model '{preferred_model}' not installed, falling back")

        # 2. Auto-pick best available by priority list
        if not vision_model:
            priority_substrings = [
                "qwen2-vl", "qwen2.5-vl", "qwen-vl",
                "llava:34b", "llava:13b", "llava-llama3",
                "llava", "bakllava", "llama-vision", "llama3.2-vision",
                "moondream",
            ]
            installed_lower = [(m, m.lower()) for m in installed]
            for needle in priority_substrings:
                for orig, low in installed_lower:
                    if needle in low:
                        vision_model = orig
                        break
                if vision_model:
                    break

        if not vision_model:
            logger.info("No vision model available in Ollama, using OCR only")
            return ""

        logger.info(f"Vision: using model {vision_model} for {filename} ({len(file_bytes)} bytes)")
        img_b64 = base64.b64encode(file_bytes).decode("utf-8")

        import time
        t0 = time.time()
        # Generous timeout — vision models on CPU can take minutes for first inference
        async with httpx.AsyncClient(timeout=httpx.Timeout(600.0, connect=10.0)) as client:
            resp = await client.post(f"{base_url}/api/chat", json={
                "model": vision_model,
                "messages": [{
                    "role": "user",
                    "content": (
                        "Подробно опиши, что изображено на картинке. "
                        "Если это техническое оборудование — назови узлы, компоненты, маркировки, цвета проводов/волокон. "
                        "Если есть любой текст, цифры, идентификаторы — перепиши их дословно."
                    ),
                    "images": [img_b64],
                }],
                "stream": False,
                "options": {"num_predict": 600, "temperature": 0.2},
            })
            resp.raise_for_status()
            data = resp.json()
            elapsed = time.time() - t0
            content = data.get("message", {}).get("content", "").strip()
            logger.info(f"Vision: model {vision_model} returned {len(content)} chars in {elapsed:.1f}s")
            return content

    except Exception as e:
        logger.warning(
            f"Vision description failed for {filename} (model={vision_model or 'unknown'}): "
            f"{type(e).__name__}: {e!r}"
        )
        return ""


async def extract_audio_whisper(file_bytes: bytes, filename: str) -> str:
    """Transcribe audio via host faster-whisper-server (Whisper large-v3 on GPU).
    Replaces the old local 'openai-whisper' base model — much more accurate for
    ru/uk and ~3-5× faster (GPU-bound)."""
    import os
    import httpx

    url = os.getenv("WHISPER_URL", "http://172.10.100.9:8001/v1/audio/transcriptions")
    model = os.getenv("WHISPER_MODEL", "Systran/faster-whisper-large-v3")

    ext = os.path.splitext(filename)[1].lower().lstrip(".") or "wav"
    mime = {
        "mp3": "audio/mpeg",
        "wav": "audio/wav",
        "ogg": "audio/ogg",
        "m4a": "audio/mp4",
        "aac": "audio/aac",
        "flac": "audio/flac",
        "opus": "audio/opus",
        "webm": "audio/webm",
        "wma": "audio/x-ms-wma",
    }.get(ext, "application/octet-stream")

    try:
        async with httpx.AsyncClient(timeout=300) as client:
            r = await client.post(
                url,
                files={
                    "file": (filename, file_bytes, mime),
                    "model": (None, model),
                    "response_format": (None, "json"),
                },
            )
            r.raise_for_status()
            data = r.json()
            text = (data.get("text") or "").strip()
            logger.info(
                f"Whisper transcribed {filename} ({len(file_bytes)} bytes) → {len(text)} chars text"
            )
            return text
    except httpx.HTTPStatusError as e:
        logger.error(
            f"Whisper HTTP {e.response.status_code} for {filename}: "
            f"{e.response.text[:300] if e.response else ''}"
        )
        raise
    except Exception as e:
        logger.error(f"Whisper transcription failed for {filename}: {e}")
        raise


async def extract_content(file_bytes: bytes, filename: str, file_type: str, vision_model: str | None = None) -> str:
    """Route to the appropriate extractor based on file type."""
    if file_type == "pdf":
        # Text-layer first, OCR fallback for scanned pages.
        return await extract_pdf_content_async(file_bytes, filename)
    elif file_type == "docx":
        return extract_docx_content(file_bytes)
    elif file_type == "xlsx":
        return extract_xlsx_content(file_bytes)
    elif file_type == "image":
        # Fast path: GPU PaddleOCR for any text on the image.
        ocr_text = await extract_image_ocr(file_bytes, filename)
        parts: list[str] = []
        if ocr_text:
            parts.append(f"[Распознанный текст (OCR)]\n{ocr_text}")
        # Optional slow path: VLM image description (only when explicitly enabled —
        # otherwise CPU vision can spin for minutes per image).
        if settings.ENABLE_CPU_VISION_DESCRIPTION:
            vision_text = await describe_image_vision(file_bytes, filename, preferred_model=vision_model)
            if vision_text:
                parts.append(f"[Описание изображения]\n{vision_text}")
        return "\n\n".join(parts) if parts else ""
    elif file_type == "audio":
        # Async HTTP call to host faster-whisper-server (Whisper large-v3 on GPU).
        return await extract_audio_whisper(file_bytes, filename)
    else:
        # text, csv, json, html, xml, md, etc.
        return extract_text_content(file_bytes, filename)


# ============================================================
# Main processing pipeline
# ============================================================

async def process_attachment(
    attachment_id: uuid.UUID,
    tenant_id: uuid.UUID,
    db: AsyncSession,
) -> None:
    """
    Process an attachment: extract text, summarize, chunk, embed.
    """
    att = (await db.execute(
        select(MessageAttachment).where(MessageAttachment.id == attachment_id)
    )).scalar_one_or_none()

    if not att:
        logger.error(f"Attachment {attachment_id} not found")
        return

    att.processing_status = "processing"
    att.processing_error = None
    await db.flush()

    try:
        # 1. Read file
        file_bytes = await read_file(att.storage_path)

        # 2. Get tenant config (used for vision model + embedding model + summary LLM)
        config = (await db.execute(
            select(TenantShellConfig).where(TenantShellConfig.tenant_id == tenant_id)
        )).scalar_one_or_none()
        if not config:
            raise ValueError("Shell config not found for tenant")

        embedding_model = config.embedding_model_name
        if not embedding_model:
            raise ValueError("Embedding model not configured")

        # 3. Extract text based on file type (uses config.vision_model_name for images)
        content_text = await extract_content(
            file_bytes, att.filename, att.file_type,
            vision_model=config.vision_model_name,
        )
        if not content_text or not content_text.strip():
            raise ValueError(f"No content could be extracted from {att.file_type} file")

        att.content_text = content_text

        # Use Ollama for embeddings (local)
        embed_provider = get_provider("ollama", settings.OLLAMA_BASE_URL or "http://localhost:11434")

        # 4. Generate summary via LLM
        from app.services.llm.model_resolver import _resolve_from_shell_config
        resolved = _resolve_from_shell_config(config)

        try:
            summary_text = content_text[:3000]
            from app.services.llm.language import build_language_pin_message
            from app.services.attachments.summary_parser import generate_attachment_summary
            att.summary = await generate_attachment_summary(
                content=summary_text,
                provider=resolved.provider,
                model_name=resolved.model_name,
                language=config.response_language,
            )
            if not att.summary:
                att.summary = f"Файл: {att.filename} ({att.file_type}, {att.file_size_bytes} байт)"
        except Exception as e:
            logger.warning(f"Summary generation failed for attachment {attachment_id}: {e}")
            att.summary = f"Файл: {att.filename} ({att.file_type}, {att.file_size_bytes} байт)"

        # 5. Chunk text
        chunks = chunk_text(content_text)
        if not chunks:
            raise ValueError("No chunks generated from content")

        # 6. Delete old chunks
        await db.execute(
            delete(MessageAttachmentChunk).where(MessageAttachmentChunk.attachment_id == attachment_id)
        )

        # 7. Generate embeddings in batches
        BATCH_SIZE = 10
        all_chunks = []

        for i in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[i:i + BATCH_SIZE]
            embeddings = await embed_provider.embed(batch, embedding_model)

            for j, (chunk_text_item, emb) in enumerate(zip(batch, embeddings)):
                chunk = MessageAttachmentChunk(
                    attachment_id=att.id,
                    tenant_id=tenant_id,
                    chunk_index=i + j,
                    content=chunk_text_item,
                    embedding=emb,
                )
                db.add(chunk)
                all_chunks.append(chunk)

        att.chunks_count = len(all_chunks)
        att.processing_status = "done"
        att.processing_error = None
        await db.flush()

        logger.info(f"Attachment {attachment_id}: {att.file_type} extracted, {len(all_chunks)} chunks embedded")

    except Exception as e:
        att.processing_status = "error"
        att.processing_error = str(e)[:500]
        await db.flush()
        logger.error(f"Attachment {attachment_id} processing failed: {e}")


async def search_attachment_chunks(
    attachment_id: str,
    query: str,
    db: AsyncSession,
    provider,
    embedding_model: str,
    max_results: int = 5,
) -> list[MessageAttachmentChunk]:
    """
    Semantic search within a specific attachment's chunks.
    """
    query_embeddings = await provider.embed(query, embedding_model)
    query_vector = query_embeddings[0]

    stmt = (
        select(MessageAttachmentChunk)
        .where(
            MessageAttachmentChunk.attachment_id == attachment_id,
            MessageAttachmentChunk.embedding.isnot(None),
        )
        .order_by(MessageAttachmentChunk.embedding.cosine_distance(query_vector))
        .limit(max_results)
    )

    result = await db.execute(stmt)
    return list(result.scalars().all())


async def process_attachment_background(attachment_id: str, tenant_id: str) -> None:
    """Background entrypoint with its own DB session and transaction."""
    async with async_session() as db:
        try:
            await process_attachment(uuid.UUID(str(attachment_id)), uuid.UUID(str(tenant_id)), db)
            await db.commit()
        except Exception:
            await db.rollback()
            logger.exception("Attachment background processing failed", extra={"attachment_id": attachment_id})
            return
    # After successful processing, promote the attachment to a first-class
    # Artifact so auto-grounding can surface it on later turns alongside
    # code/script artifacts. Idempotent — re-processing updates in place.
    try:
        from app.services.artifacts.from_attachment import upsert_artifact_from_attachment
        await upsert_artifact_from_attachment(attachment_id=uuid.UUID(str(attachment_id)))
    except Exception:
        logger.exception("Attachment → artifact upsert failed (non-fatal)", extra={"attachment_id": attachment_id})
